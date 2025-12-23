# rag_app.py - Cognitbotz AI Legal Platform with Authentication + Legal Drafting + Auto-Embedding

import streamlit as st
import time
import os
import tempfile
import sys
from pathlib import Path
from dotenv import load_dotenv, find_dotenv
import sqlite3
from datetime import datetime
import uuid
import logging
import atexit
import glob

# Load environment first
load_dotenv(find_dotenv())

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from vectors import EmbeddingsManager  
from chatbot import ChatbotManager
from legal_drafting import LegalDraftingManager

# PDF and DOCX generation libraries
from io import BytesIO
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Suppress warnings
import warnings
warnings.filterwarnings('ignore', message='.*Qdrant client version.*incompatible.*')

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = ['pdf', 'txt', 'docx', 'pptx']
MAX_INTERACTIONS_PER_SESSION = 25
MAX_TOKENS_PER_SESSION = 75000

# PERSISTENT COLLECTION NAME - Same as vectors.py and chatbot.py
PERSISTENT_COLLECTION_NAME = "Legal_documents"

# Path to static documents folder
DOCUMENTS_FOLDER = Path(__file__).parent / "documents"

DEFAULT_CONFIG = {
    'llm_provider': 'Groq',
    'llm_model': 'llama-3.3-70b-versatile',
    'temperature': 0.3,
    'max_tokens': 3000,
    'retrieval_k': 5,
    'score_threshold': 0.5,
    'enable_content_filter': True,
    'enable_pii_detection': True,
    'max_interactions': 25,
    'chatbot_mode': 'Document Only'
}

# ==================== AUTO-EMBEDDING FUNCTIONALITY ====================

def check_collection_exists():
    """Check if the persistent collection exists in Qdrant"""
    try:
        from qdrant_client import QdrantClient
        
        qdrant_url = os.getenv('QDRANT_URL')
        qdrant_api_key = os.getenv('QDRANT_API_KEY')
        
        if not qdrant_url or not qdrant_api_key:
            logger.warning("Qdrant credentials not found")
            return False
        
        client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key, prefer_grpc=False)
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]
        
        exists = PERSISTENT_COLLECTION_NAME in collection_names
        
        if exists:
            # Check if collection has any documents
            try:
                collection_info = client.get_collection(PERSISTENT_COLLECTION_NAME)
                points_count = collection_info.points_count
                logger.info(f"Collection '{PERSISTENT_COLLECTION_NAME}' exists with {points_count} embeddings")
                return points_count > 0
            except Exception as e:
                logger.error(f"Error checking collection info: {e}")
                return False
        
        return False
        
    except Exception as e:
        logger.error(f"Error checking collection: {e}")
        return False

def get_documents_from_folder():
    """Get all documents from the static documents folder"""
    if not DOCUMENTS_FOLDER.exists():
        logger.warning(f"Documents folder not found: {DOCUMENTS_FOLDER}")
        return []
    
    document_files = []
    for ext in ALLOWED_EXTENSIONS:
        pattern = str(DOCUMENTS_FOLDER / f"*.{ext}")
        files = glob.glob(pattern)
        document_files.extend(files)
    
    logger.info(f"Found {len(document_files)} documents in folder")
    return document_files

def auto_embed_documents():
    """Automatically embed documents from the static folder on first run"""
    try:
        # Check if embeddings already exist
        if check_collection_exists():
            logger.info("Embeddings already exist. Skipping auto-embedding.")
            return True, "Embeddings loaded successfully"
        
        # Get documents from folder
        document_files = get_documents_from_folder()
        
        if not document_files:
            logger.warning("No documents found in documents folder")
            return False, "No documents found to embed"
        
        # Create embeddings
        logger.info(f"Creating embeddings for {len(document_files)} documents...")
        
        embeddings_mgr = EmbeddingsManager(
            model_name="BAAI/bge-small-en",
            device="cpu",
            encode_kwargs={"normalize_embeddings": True},
            qdrant_url=os.getenv('QDRANT_URL'),
            collection_name=PERSISTENT_COLLECTION_NAME,
            chunk_size=1200,
            chunk_overlap=300
        )
        
        processed_count = 0
        failed_count = 0
        
        for doc_path in document_files:
            try:
                logger.info(f"Processing: {Path(doc_path).name}")
                embeddings_mgr.create_embeddings(doc_path)
                processed_count += 1
            except Exception as e:
                logger.error(f"Failed to process {doc_path}: {e}")
                failed_count += 1
        
        if processed_count > 0:
            message = f"✅ Auto-embedded {processed_count} documents. Chat enabled!"
            logger.info(message)
            return True, message
        else:
            return False, "Failed to embed documents"
            
    except Exception as e:
        logger.error(f"Auto-embedding error: {e}")
        return False, f"Error: {str(e)}"

# ==================== PDF AND DOCX GENERATION ====================

def generate_pdf(content: str, doc_type: str, metadata: dict) -> BytesIO:
    """Generate professional PDF from document content"""
    if not REPORTLAB_AVAILABLE:
        return None
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    
    # Custom styles for legal documents
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='#1e3a8a',
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='#1e3a8a',
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=12,
        fontName='Times-Roman',
        leading=14
    )
    
    story = []
    
    # Add title
    story.append(Paragraph(doc_type.upper(), title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Add metadata
    metadata_text = f"<b>Generated:</b> {metadata.get('generated_at', 'N/A')}<br/>"
    metadata_text += f"<b>Jurisdiction:</b> {metadata.get('jurisdiction', 'General')}"
    story.append(Paragraph(metadata_text, body_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Process content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1*inch))
            continue
        
        # Detect headings (all caps or starts with numbers)
        if line.isupper() and len(line) < 100:
            story.append(Paragraph(line, heading_style))
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', 'Article', 'Section')):
            story.append(Paragraph(f"<b>{line}</b>", body_style))
        else:
            story.append(Paragraph(line, body_style))
    
    # Add Cognitbotz footer
    story.append(Spacer(1, 0.5*inch))
    footer_text = "<i>This document was generated by Cognitbotz AI Legal Platform. " \
                 "Please review with qualified legal counsel before use.</i>"
    story.append(Paragraph(footer_text, ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor='gray',
        alignment=TA_CENTER
    )))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx(content: str, doc_type: str, metadata: dict) -> BytesIO:
    """Generate professional DOCX from document content"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = DocxDocument()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = doc_type
    core_props.author = "Cognitbotz AI Legal Platform"
    core_props.comments = f"Generated on {metadata.get('generated_at', 'N/A')}"
    
    # Add title
    title = doc.add_heading(doc_type.upper(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.add_run('Generated: ').bold = True
    meta_para.add_run(metadata.get('generated_at', 'N/A'))
    meta_para.add_run('\nJurisdiction: ').bold = True
    meta_para.add_run(metadata.get('jurisdiction', 'General'))
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    # Process content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Detect headings
        if line.isupper() and len(line) < 100:
            heading = doc.add_heading(line, level=2)
            heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', 'Article', 'Section')):
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
        else:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add Cognitbotz footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        'This document was generated by Cognitbotz AI Legal Platform. '
        'Please review with qualified legal counsel before use.'
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== ENVIRONMENT VALIDATION ====================

def validate_environment():
    """Validate required environment variables at startup"""
    required_vars = {
        'QDRANT_URL': 'Qdrant vector database URL',
        'QDRANT_API_KEY': 'Qdrant API key',
        'GROQ_API_KEY': 'Groq API key for LLM'
    }
    
    missing_vars = []
    for var, description in required_vars.items():
        if not os.getenv(var):
            missing_vars.append(f"- {var}: {description}")
    
    if missing_vars:
        return False, missing_vars
    
    return True, []

# ==================== AUTHENTICATION ====================

# Hardcoded users with plain passwords (simple approach)
HARDCODED_USERS = {
    "admin": {
        "password": "Admin@123",
        "role": "Administrator",
        "full_name": "System Administrator"
    },
    "finance": {
        "password": "Finance@123",
        "role": "Finance User",
        "full_name": "Finance Department"
    },
    "requester": {
        "password": "Request@123",
        "role": "Purchase Requester",
        "full_name": "Purchase Requester"
    }
}

def login_page():
    """Render login page"""
    
    # Center the login form
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        
        st.title("⚖️ Cognitbotz AI Legal Platform")
        st.subheader("Please login to continue")
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Login form
        username = st.text_input(
            "Username",
            key="login_username",
            placeholder="Enter your username"
        )
        
        password = st.text_input(
            "Password",
            type="password",
            key="login_password",
            placeholder="Enter your password"
        )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Login button
        if st.button("🔐 Login", use_container_width=True, type="primary"):
            if not username or not password:
                st.error("⚠️ Please enter both username and password")
            elif username not in HARDCODED_USERS:
                st.error("❌ Invalid username or password")
            else:
                # Simple password check (no hashing)
                user_data = HARDCODED_USERS[username]
                
                if password == user_data["password"]:
                    # Authentication successful
                    st.session_state.authenticated = True
                    st.session_state.username = username
                    st.session_state.user_role = user_data["role"]
                    st.session_state.full_name = user_data["full_name"]
                    st.session_state.login_time = datetime.now()
                    
                    st.success("✅ Login successful! Redirecting...")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ Invalid username or password")
        
        # Test credentials info
        with st.expander("🔑 Test Credentials"):
            st.markdown("""
            **Administrator:**
            - Username: `admin`
            - Password: `Admin@123`
            
            **Finance User:**
            - Username: `finance`
            - Password: `Finance@123`
            
            **Purchase Requester:**
            - Username: `requester`
            - Password: `Request@123`
            """)

# ==================== DATABASE FUNCTIONS ====================

def validate_file_upload(uploaded_file):
    """Validate file uploads"""
    if uploaded_file is None:
        return False, "No file uploaded"
    
    if uploaded_file.size > MAX_FILE_SIZE:
        return False, f"File too large. Max: {MAX_FILE_SIZE // (1024*1024)}MB"
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        return False, f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
    
    if any(char in uploaded_file.name for char in ['<', '>', '..', '/', '\\']):
        return False, "Invalid characters in filename"
    
    return True, "Valid file"

def migrate_database():
    """Ensure database has correct schema"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        cursor.execute("PRAGMA table_info(chat_logs)")
        existing_columns = [row[1] for row in cursor.fetchall()]
        
        if not existing_columns:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chat_logs (
                    id TEXT PRIMARY KEY,
                    username TEXT,
                    chat_session_id TEXT,
                    session_id TEXT,
                    timestamp TEXT,
                    user_input TEXT,
                    bot_response TEXT,
                    tokens_used INTEGER,
                    is_flagged BOOLEAN,
                    flag_reason TEXT
                )
            """)
        else:
            if 'username' not in existing_columns:
                cursor.execute("ALTER TABLE chat_logs ADD COLUMN username TEXT")
            if 'chat_session_id' not in existing_columns:
                cursor.execute("ALTER TABLE chat_logs ADD COLUMN chat_session_id TEXT")
        
        conn.commit()
        conn.close()
        logger.info("Database migration successful")
        
    except Exception as e:
        logger.error(f"Database migration error: {e}")

def log_interaction(session_id, user_input, bot_response, tokens_used, is_flagged=False, flag_reason=None, username=None, chat_session_id=None):
    """Log chat interaction to database"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        log_id = str(uuid.uuid4())
        timestamp = datetime.now().isoformat()
        
        cursor.execute("""
            INSERT INTO chat_logs 
            (id, username, chat_session_id, session_id, timestamp, user_input, bot_response, tokens_used, is_flagged, flag_reason)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (log_id, username, chat_session_id, session_id, timestamp, user_input, bot_response, tokens_used, is_flagged, flag_reason))
        
        conn.commit()
        conn.close()
        
    except Exception as e:
        logger.error(f"Failed to log interaction: {e}")

def load_user_chat_sessions(username):
    """Load all chat sessions for a user"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT 
                chat_session_id,
                MIN(timestamp) as first_msg_time,
                COUNT(*) as message_count,
                user_input
            FROM chat_logs
            WHERE username = ? AND chat_session_id IS NOT NULL
            GROUP BY chat_session_id
            ORDER BY first_msg_time DESC
        """, (username,))
        
        rows = cursor.fetchall()
        conn.close()
        
        sessions = []
        for row in rows:
            chat_id, timestamp, count, first_input = row
            
            try:
                dt = datetime.fromisoformat(timestamp)
                date_str = dt.strftime("%b %d, %Y %I:%M %p")
            except:
                date_str = timestamp
            
            preview = (first_input[:40] + "...") if first_input and len(first_input) > 40 else (first_input or "Empty chat")
            
            sessions.append({
                "id": chat_id,
                "date": date_str,
                "count": count,
                "preview": preview
            })
        
        return sessions
        
    except Exception as e:
        logger.error(f"Error loading chat sessions: {e}")
        return []

def load_chat_messages(chat_session_id):
    """Load all messages from a specific chat session"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT user_input, bot_response, tokens_used, is_flagged, flag_reason, timestamp
            FROM chat_logs
            WHERE chat_session_id = ?
            ORDER BY timestamp ASC
        """, (chat_session_id,))
        
        rows = cursor.fetchall()
        conn.close()
        
        messages = []
        for row in rows:
            user_input, bot_response, tokens, flagged, flag_reason, _ = row
            
            messages.append({
                "role": "user",
                "content": user_input
            })
            messages.append({
                "role": "assistant",
                "content": bot_response,
                "tokens_used": tokens,
                "is_flagged": bool(flagged),
                "flag_reason": flag_reason,
                "sources": []
            })
        
        return messages
        
    except Exception as e:
        logger.error(f"Error loading chat messages: {e}")
        return []

def delete_chat_session(chat_session_id, username):
    """Delete a chat session and all its messages"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        cursor.execute("""
            DELETE FROM chat_logs 
            WHERE chat_session_id = ? AND username = ?
        """, (chat_session_id, username))
        
        conn.commit()
        deleted_count = cursor.rowcount
        conn.close()
        
        return deleted_count > 0
        
    except Exception as e:
        logger.error(f"Error deleting chat session: {e}")
        return False

def start_new_chat():
    """Start a brand new chat"""
    st.session_state.chat_session_id = str(uuid.uuid4())
    st.session_state.messages = []
    st.session_state.interaction_count = 0
    st.session_state.total_tokens_used = 0
    logger.info(f"Started new chat: {st.session_state.chat_session_id}")

# ==================== PAGE CONFIG ====================

st.set_page_config(
    page_title="Cognitbotz AI Legal Platform",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .chat-message {
        padding: 1rem;
        margin: 0.5rem 0; 
        border-radius: 0.5rem; 
        border-left: 4px solid #059669;
        background-color: #f0fdf4;
    }
    .user-message { 
        background-color: #eff6ff; 
        border-left-color: #2563eb; 
        margin-left: 10%;
    }
    .assistant-message { 
        background-color: #f9fafb; 
        border-left-color: #059669; 
        margin-right: 10%;
    }
    .source-info {
        background-color: #f8f9fa;
        padding: 0.5rem;
        border-radius: 0.3rem;
        margin: 0.2rem 0;
        font-size: 0.85rem;
        border-left: 3px solid #059669;
    }
    .response-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 0.25rem;
        font-size: 0.75rem;
        font-weight: 600;
        margin-bottom: 0.5rem;
    }
    .badge-rag {
        background-color: #d4edda;
        color: #155724;
        border: 1px solid #c3e6cb;
    }
    .badge-general {
        background-color: #fff3cd;
        color: #856404;
        border: 1px solid #ffeaa7;
    }
    .badge-drafting {
        background-color: #cfe2ff;
        color: #084298;
        border: 1px solid #b6d4fe;
    }
    .footer-disclaimer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background-color: #fff;
        border-top: 2px solid #f0f0f0;
        padding: 0.75rem 1rem;
        text-align: center;
        color: #666;
        font-size: 0.85rem;
        z-index: 999;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
    }
    .footer-disclaimer p {
        margin: 0;
        padding: 0;
    }
    .stChatInputContainer {
        padding-bottom: 60px;
    }
    .doc-preview {
        background-color: #ffffff;
        border: 2px solid #e5e7eb;
        border-radius: 0.75rem;
        padding: 2rem;
        margin: 1rem 0;
        max-height: 600px;
        overflow-y: auto;
        font-family: 'Times New Roman', serif;
        line-height: 1.6;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .doc-preview h1, .doc-preview h2, .doc-preview h3 {
        color: #1e3a8a;
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
    }
    .doc-preview p {
        text-align: justify;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state"""
    defaults = {
        'authenticated': False,
        'temp_file_paths': [],
        'session_id': str(uuid.uuid4()),
        'chat_session_id': str(uuid.uuid4()),
        'chat_sessions': [],
        'files_processed': False,
        'embeddings_saved': False,
        'chatbot_manager': None,
        'drafting_manager': None,
        'last_generated_doc': None,
        'drafted_documents': [],
        'messages': [],
        'document_info': {},
        'interaction_count': 0,
        'total_tokens_used': 0,
        'last_query_tokens': 0,
        'input_tokens_used': 0,
        'output_tokens_used': 0,
        'error_count': 0,
        'processed_files': [],
        'processing_status': {},
        'collection_created': False,
        'processing': False,
        'token_warnings_shown': [],
        'db_migrated': False,
        'env_validated': False,
        'user_preferences': {},
        'auto_embedding_done': False,
        'auto_embedding_checked': False
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

# ==================== MAIN APPLICATION ====================

# Initialize
initialize_session_state()

# Handle redirects at the very beginning to avoid widget conflicts
if "redirect_to_draft" in st.session_state:
    st.session_state.chatbot_mode = "Legal Drafting"
    del st.session_state["redirect_to_draft"]
    st.rerun()

if "redirect_to_chat" in st.session_state:
    st.session_state.chatbot_mode = "General Chat"
    del st.session_state["redirect_to_chat"]
    st.rerun()

# Validate environment variables
if not st.session_state.get('env_validated'):
    env_valid, missing = validate_environment()
    
    if not env_valid:
        st.error("⚠️ **Environment Configuration Error**")
        st.markdown("The following environment variables are missing or not set:")
        for var in missing:
            st.markdown(f"  {var}")
        st.markdown("---")
        st.info("""
        **How to fix:**
        1. Create a `.env` file in your project root
        2. Add the following variables:
        ```
        QDRANT_URL=your_qdrant_url
        QDRANT_API_KEY=your_qdrant_api_key
        GROQ_API_KEY=your_groq_api_key
        ```
        3. Restart the application
        """)
        st.stop()
    
    st.session_state.env_validated = True

# Run database migrations
if not st.session_state.get('db_migrated'):
    migrate_database()
    st.session_state.db_migrated = True

# Check authentication
if not st.session_state.get('authenticated'):
    login_page()
    st.stop()

# AUTO-EMBED DOCUMENTS ON FIRST RUN
if not st.session_state.get('auto_embedding_checked'):
    with st.spinner("🔍 Checking for existing embeddings..."):
        # First check if embeddings exist
        embeddings_exist = check_collection_exists()
        
        if embeddings_exist:
            st.session_state['collection_created'] = True
            st.session_state['auto_embedding_done'] = True
            st.success("✅ Embeddings loaded! You can now chat with the legal documents.")
            logger.info("Embeddings already exist - chat enabled")
        else:
            # Try to auto-embed from documents folder
            with st.status("📚 Processing legal documents from library...", expanded=True) as status:
                st.write("Checking documents folder...")
                
                success, message = auto_embed_documents()
                
                if success:
                    st.session_state['collection_created'] = True
                    st.session_state['auto_embedding_done'] = True
                    status.update(label=message, state="complete")
                    st.success(message)
                else:
                    status.update(label=f"⚠️ {message}", state="error")
                    st.warning(f"⚠️ {message}")
                    st.info("You can still upload documents manually using the sidebar.")
        
        st.session_state['auto_embedding_checked'] = True
        time.sleep(2)

# Load chat history for authenticated user
if not st.session_state.chat_sessions:
    st.session_state.chat_sessions = load_user_chat_sessions(st.session_state.username)

# ==================== SIDEBAR ====================
with st.sidebar:
    st.markdown("### ⚖️ Cognitbotz AI Legal Platform")
    st.markdown(f"**Welcome, {st.session_state.get('full_name', st.session_state.username)}!**")
    
    st.markdown("---")
    
    # Chat History
    st.markdown("### 💬 Chat History")
    
    if st.button("➕ New Chat", type="primary", use_container_width=True, key="new_chat_btn"):
        start_new_chat()
        st.rerun()
    
    chat_sessions = load_user_chat_sessions(st.session_state.username)
    
    if chat_sessions:
        st.markdown(f"**Previous Chats ({len(chat_sessions)})**")
        
        for i, session in enumerate(chat_sessions[:10]):
            col1, col2 = st.columns([4, 1])
            
            with col1:
                button_label = f"💬 {session['preview']}"
                if st.button(button_label, key=f"chat_{i}_{session['id']}", use_container_width=True):
                    st.session_state.chat_session_id = session['id']
                    st.session_state.messages = load_chat_messages(session['id'])
                    st.session_state.interaction_count = session['count'] // 2
                    
                    # FIXED: Ensure collection is marked as created when loading previous chat
                    if st.session_state.get('auto_embedding_done') or check_collection_exists():
                        st.session_state['collection_created'] = True
                    
                    st.success(f"✅ Loaded: {session['preview']}")
                    st.rerun()
            
            with col2:
                if st.button("🗑️", key=f"del_{i}_{session['id']}", help="Delete this chat"):
                    if delete_chat_session(session['id'], st.session_state.username):
                        st.success("✅ Deleted!")
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error("❌ Failed")
            
            st.caption(f"📅 {session['date']} • {session['count']} msgs")
            st.markdown("---")
    else:
        st.info("No previous chats")
    
    st.markdown("---")
    
        
    # Chatbot Mode
    st.markdown("### 🤖 Mode Selection")
    chatbot_mode = st.selectbox(
        "Select Mode",
        ["Document Only", "General Chat", "Hybrid (Smart)", "Layman Explanation"],
        index=0,
        key="chatbot_mode"
    )
    
    st.markdown("---")
    
    # Session Statistics
    st.markdown("### 📊 Session Statistics")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Interactions", f"{st.session_state.interaction_count}/{DEFAULT_CONFIG['max_interactions']}")
        
        # Show if embeddings are ready
        if st.session_state.get('collection_created'):
            st.metric("📚 Embeddings", "✅ Ready")
        else:
            st.metric("📚 Embeddings", "⏳ Pending")
    with col2:
        tokens = st.session_state.get('total_tokens_used', 0)
        st.metric("Tokens", f"{tokens:,}")
        st.metric("Remaining", f"{MAX_TOKENS_PER_SESSION - tokens:,}")
    
    st.progress(min(tokens / MAX_TOKENS_PER_SESSION, 1.0))
    
    st.markdown("---")
    
    # User Info
    st.markdown("### 👤 User Info")
    st.write(f"**User:** {st.session_state.get('full_name', st.session_state.username)}")
    st.write(f"**Role:** {st.session_state.get('user_role', 'User')}")
    
    if st.session_state.get('login_time'):
        duration = datetime.now() - st.session_state['login_time']
        st.write(f"**Session:** {int(duration.total_seconds() / 60)} min")
    
    st.markdown("---")
    
    # Session Management
    st.markdown("### 🔄 Session Management")
    
    if st.button("🔄 Reset Session", type="secondary", use_container_width=True):
        if st.session_state.get('chatbot_manager'):
            try:
                st.session_state['chatbot_manager'].cleanup_collection()
            except:
                pass
        
        keys_to_keep = ['authenticated', 'username', 'full_name', 'user_role', 'user_email', 'login_time', 'db_migrated', 'env_validated', 'auto_embedding_done', 'auto_embedding_checked', 'collection_created']
        keys_to_delete = [key for key in st.session_state.keys() if key not in keys_to_keep]
        
        for key in keys_to_delete:
            del st.session_state[key]
        
        initialize_session_state()
        # Restore the embedding status
        st.session_state['auto_embedding_checked'] = True
        if check_collection_exists():
            st.session_state['collection_created'] = True
            st.session_state['auto_embedding_done'] = True
        
        st.success("✅ Session reset!")
        time.sleep(1)
        st.rerun()
    
    st.markdown("---")
    
    # Logout
    if st.button("🚪 Logout", use_container_width=True, type="secondary"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        
        st.rerun()

# ==================== MAIN CONTENT ====================

# Check if Legal Drafting mode is selected
if chatbot_mode == "Legal Drafting":
    # ==================== LEGAL DRAFTING MODE ====================
    st.title("✏️ Legal Document Drafting")
    
    # Add ASK and DRAFT buttons below the title
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("💬 ASK", use_container_width=True, type="secondary"):
            st.session_state.main_mode = "ASK"
            st.rerun()
    
    with col2:
        if st.button("✍️ DRAFT", use_container_width=True, type="primary"):
            st.session_state.main_mode = "DRAFT"
            st.rerun()
    
    # Handle mode changes
    main_mode = st.session_state.get("main_mode", "DRAFT")
    chatbot_mode = st.session_state.get("chatbot_mode", "Legal Drafting")
    
    if main_mode == "ASK" and chatbot_mode != "General Chat":
        # Redirect to chat mode
        st.session_state.redirect_to_chat = True
        st.rerun()
    elif main_mode == "DRAFT" and chatbot_mode == "General Chat":
        # Redirect to drafting mode
        st.session_state.redirect_to_draft = True
        st.rerun()
    
    st.markdown("Generate professional legal documents with AI assistance")
    
    # Initialize drafting manager
    if not st.session_state.get('drafting_manager'):
        try:
            drafting_manager = LegalDraftingManager(
                llm_model=DEFAULT_CONFIG['llm_model'],
                temperature=0.3,
                max_tokens=12000,
                user_preferences=st.session_state.get('user_preferences', {})
            )
            st.session_state['drafting_manager'] = drafting_manager
        except Exception as e:
            st.error(f"❌ Failed to initialize drafting: {e}")
            st.stop()
    
    # Document configuration
    col1, col2 = st.columns(2)
    
    with col1:
        doc_type = st.selectbox(
            "Document Type",
            [
                "Contracts & Agreements",
                "Petitions & Applications",
                "Court Orders & Judgments",
                "Legal Briefs & Submissions",
                "Statutes & Regulations"
            ]
        )
        
        draft_style = st.selectbox(
            "Drafting Style",
            ["Formal Legal", "Plain English", "Business Formal", "Technical", "Academic Legal"],
            help="Choose the writing style"
        )
    
    with col2:
        document_length = st.selectbox(
            "Document Length",
            ["Brief", "Short (1-2 pages)", "Standard", "Medium (3-5 pages)", "Comprehensive", "Long (6-10 pages)", "Detailed"],
            help="Select desired document length"
        )
        
        # Optional clauses
        available_clauses = [
            "Definitions", "Parties", "Recitals", "Terms and Conditions",
            "Representations", "Warranties", "Limitations", "Indemnification",
            "Confidentiality", "Termination", "Dispute Resolution", "Governing Law"
        ]
        selected_clauses = st.multiselect("Include Clauses (optional)", available_clauses)
    
    # Special provisions
    special_provisions = st.text_area(
        "Special Provisions (Optional):",
        placeholder="Any special terms, conditions, or provisions...",
        height=100,
        help="Add specific requirements or provisions"
    )
    
    # Main drafting requirements
    st.markdown("### 📝 Document Requirements")
    requirements = st.text_area(
        "Describe what you want to draft:",
        height=200,
        placeholder="Example: Draft a service agreement between Company A and Company B for software development services. Include payment terms of $10,000 per month, 6-month initial term, confidentiality obligations...",
        help="Provide detailed requirements"
    )
    
    # Generate and clear buttons
    col1, col2, col3 = st.columns([1, 1, 2])
    
    with col1:
        generate_btn = st.button("✨ Generate Document", type="primary", use_container_width=True)
    
    with col2:
        if st.session_state.get('last_generated_doc'):
            if st.button("🗑️ Clear Document", type="secondary", use_container_width=True):
                st.session_state['last_generated_doc'] = None
                st.rerun()
    
    # Generate document
    if generate_btn:
        if not requirements.strip():
            st.warning("⚠️ Please provide document requirements")
        else:
            with st.spinner("✏️ Drafting your legal document..."):
                try:
                    result = st.session_state['drafting_manager'].generate_document(
                        doc_type=doc_type,
                        prompt=requirements,
                        style=draft_style,
                        length=document_length,
                        clauses=selected_clauses if selected_clauses else None,
                        special_provisions=special_provisions if special_provisions else ""
                    )
                    
                    if 'error' not in result:
                        st.session_state['last_generated_doc'] = result
                        st.session_state.drafted_documents.append(result)
                        st.session_state.total_tokens_used += result.get('tokens_used', 0)
                        
                        log_interaction(
                            st.session_state.session_id,
                            requirements,
                            result.get('document', '')[:1000],
                            result.get('tokens_used', 0),
                            username=st.session_state.get('username'),
                            chat_session_id=st.session_state.get('chat_session_id')
                        )
                        
                        st.success(f"✅ Document generated! ({result.get('word_count', 0)} words, {result.get('tokens_used', 0)} tokens)")
                        st.rerun()
                    else:
                        st.error(f"❌ Generation failed: {result['error']}")
                
                except Exception as e:
                    logger.error(f"Drafting error: {e}")
                    st.error(f"❌ Error: {str(e)}")
    
    # Display generated document
    if st.session_state.get('last_generated_doc'):
        draft_result = st.session_state['last_generated_doc']
        
        st.markdown("---")
        st.markdown("### 📄 Generated Legal Document")
        
        # Document metadata
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📋 Type", draft_result.get('doc_type', 'N/A'))
        with col2:
            st.metric("✍️ Style", draft_result.get('style', 'N/A'))
        with col3:
            st.metric("📝 Words", draft_result.get('word_count', 0))
        with col4:
            st.metric("🎯 Tokens", draft_result.get('tokens_used', 0))
        
        # Document preview
        st.markdown('<div class="doc-preview">', unsafe_allow_html=True)
        doc_content = draft_result.get('document', 'No document generated')
        
        # Better formatting: Split into lines and handle sections properly
        lines = doc_content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append('<br>')
            elif line.isupper() and len(line) < 100:
                formatted_lines.append(f'<h3 style="color: #1e3a8a; margin-top: 1.5rem;">{line}</h3>')
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', 'Article', 'Section', 'Clause')):
                formatted_lines.append(f'<p style="margin-top: 1rem;"><strong>{line}</strong></p>')
            else:
                formatted_lines.append(f'<p style="text-align: justify; margin-bottom: 0.5rem;">{line}</p>')
        
        formatted_content = '\n'.join(formatted_lines)
        st.markdown(formatted_content, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Download section
        st.markdown("### 📥 Download Options")
        
        col1, col2, col3, col4 = st.columns(4)
        
        # TXT download
        with col1:
            st.download_button(
                label="📄 Download TXT",
                data=draft_result['document'],
                file_name=f"{doc_type.lower().replace(' ', '_').replace('&', 'and')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        # PDF download
        with col2:
            if REPORTLAB_AVAILABLE:
                pdf_buffer = generate_pdf(
                    draft_result['document'],
                    draft_result['doc_type'],
                    draft_result.get('metadata', {})
                )
                if pdf_buffer:
                    st.download_button(
                        label="📕 Download PDF",
                        data=pdf_buffer,
                        file_name=f"{doc_type.lower().replace(' ', '_').replace('&', 'and')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
            else:
                st.button("📕 PDF (Install reportlab)", disabled=True, use_container_width=True)
        
        # DOCX download
        with col3:
            if DOCX_AVAILABLE:
                docx_buffer = generate_docx(
                    draft_result['document'],
                    draft_result['doc_type'],
                    draft_result.get('metadata', {})
                )
                if docx_buffer:
                    st.download_button(
                        label="📘 Download DOCX",
                        data=docx_buffer,
                        file_name=f"{doc_type.lower().replace(' ', '_').replace('&', 'and')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.button("📘 DOCX (Install python-docx)", disabled=True, use_container_width=True)
        
        # Installation instructions
        if not REPORTLAB_AVAILABLE or not DOCX_AVAILABLE:
            with st.expander("ℹ️ Enable PDF/DOCX Downloads"):
                st.markdown("""
                To enable PDF and DOCX downloads, install:
                
                ```bash
                pip install reportlab python-docx
                ```
                
                Then restart the application.
                """)
        
        # Document metadata details
        with st.expander("ℹ️ Document Metadata & Details"):
            metadata = draft_result.get('metadata', {})
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Generation Details:**")
                st.json({
                    "Generated At": metadata.get('generated_at', 'N/A'),
                    "Jurisdiction": metadata.get('jurisdiction', 'General'),
                    "Template Used": metadata.get('template_used', 'N/A')
                })
            
            with col2:
                st.markdown("**Document Statistics:**")
                st.json({
                    "Word Count": draft_result.get('word_count', 0),
                    "Input Tokens": draft_result.get('input_tokens', 0),
                    "Output Tokens": draft_result.get('output_tokens', 0),
                    "Total Tokens": draft_result.get('tokens_used', 0),
                    "Processing Time": f"{draft_result.get('processing_time', 0):.2f}s"
                })
            
            if draft_result.get('clauses_included'):
                st.markdown("**Clauses Included:**")
                st.write(", ".join(draft_result['clauses_included']))
    
    else:
        st.info("📝 No documents generated yet. Fill in the requirements above and click 'Generate Document'")
        
        with st.expander("💡 Tips for Better Results"):
            st.markdown("""
            **For Best Results:**
            - Be specific about parties involved
            - Include key terms and conditions
            - Mention jurisdiction if applicable
            - Specify any special clauses needed
            - Provide context and background
            
            **Example Requirement:**
            *"Draft a service agreement between Acme Corp (service provider) and Beta Inc (client) 
            for cloud infrastructure services. Monthly fee of $5,000, 12-month term with 30-day 
            termination notice, includes SLA guarantees of 99.9% uptime, confidentiality obligations, 
            and limitation of liability capped at annual fees. Governed by California law."*
            """)

else:
    # ==================== MAIN CHAT MODE ====================
    st.title("⚖️ Cognitbotz AI Legal Platform")
    
    # Add ASK and DRAFT buttons below the title
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("💬 ASK", use_container_width=True, type="primary"):
            st.session_state.main_mode = "ASK"
            st.rerun()
    
    with col2:
        if st.button("✍️ DRAFT", use_container_width=True, type="primary"):
            st.session_state.main_mode = "DRAFT"
            st.rerun()
    
    # Handle mode changes
    main_mode = st.session_state.get("main_mode", "ASK")
    chatbot_mode = st.session_state.get("chatbot_mode", "General Chat")
    
    if main_mode == "DRAFT" and chatbot_mode != "Legal Drafting":
        # Redirect to legal drafting mode by updating session state and rerunning
        st.session_state.redirect_to_draft = True
        st.rerun()
    elif main_mode == "ASK" and chatbot_mode == "Legal Drafting":
        # Redirect to chat mode
        st.session_state.redirect_to_chat = True
        st.rerun()
    
    st.markdown("Ask questions about your legal documents")
    
    # Check if embeddings are ready
    if not st.session_state.get('collection_created'):
        st.warning("⚠️ **Embeddings Not Ready**")
        st.info("📁 Please wait for initial document processing to complete, or upload documents in the sidebar.")
        st.markdown("**Supported formats:** PDF, TXT, DOCX, PPTX")
    else:
        # Show ready message if no messages yet
        if not st.session_state.get('messages'):
            st.success("✅ **Embeddings Ready!** You can now chat with the legal documents.")
    
    # 25-limit warning
    if st.session_state.interaction_count >= DEFAULT_CONFIG['max_interactions']:
        st.error("⚠️ **Interaction Limit Reached!**")
        st.info("💡 Click '**➕ New Chat**' in the sidebar to continue.")
        st.markdown("---")
    
    # Display messages
    if st.session_state.get('messages'):
        for msg in st.session_state['messages']:
            if msg['role'] == 'user':
                st.markdown(f"""
                <div class="chat-message user-message">
                <strong>You:</strong> {msg['content']}
                </div>
                """, unsafe_allow_html=True)
            else:
                response_type = msg.get('response_type', 'rag')
                badge_class = 'badge-rag' if response_type == 'rag' else 'badge-general'
                badge_text = '⚖️ Document' if response_type == 'rag' else '🤖 General'
                
                st.markdown(f"""
                <div class="chat-message assistant-message">
                <div class="response-badge {badge_class}">{badge_text}</div>
                <strong>Assistant:</strong> {msg['content']}
                </div>
                """, unsafe_allow_html=True)
                
                if msg.get('is_flagged'):
                    st.warning(f"⚠️ {msg.get('flag_reason', 'Flagged')}")
                
                if msg.get('tokens_used'):
                    st.caption(f"🔢 ~{msg['tokens_used']:,} tokens")
                
                if msg.get('sources') and len(msg.get('sources', [])) > 0:
                    with st.expander(f"⚖️ Sources ({len(msg['sources'])})"):
                        for i, src in enumerate(msg['sources'], 1):
                            st.markdown(f"""
                            <div class="source-info">
                            <strong>Source {i}:</strong> {src.get('file_name', 'Unknown')} 
                            (Page {src.get('page', 'N/A')})
                            </div>
                            """, unsafe_allow_html=True)
    
    # Chat input
    tokens = st.session_state.get('total_tokens_used', 0)
    can_chat = (
        st.session_state.interaction_count < DEFAULT_CONFIG['max_interactions'] and 
        tokens < MAX_TOKENS_PER_SESSION and
        (st.session_state.get('collection_created') or chatbot_mode in ["General Chat", "Hybrid (Smart)"])
    )
    
    if can_chat:
        user_input = st.chat_input("Ask your question...")
        
        if user_input:
            st.session_state.interaction_count += 1
            
            with st.spinner("🤔 Thinking..."):
                try:
                    # Initialize chatbot with persistent collection
                    if not st.session_state.get('chatbot_manager'):
                        chatbot = ChatbotManager(
                            model_name="BAAI/bge-small-en",
                            device="cpu",
                            encode_kwargs={"normalize_embeddings": True},
                            llm_model=DEFAULT_CONFIG['llm_model'],
                            llm_temperature=DEFAULT_CONFIG['temperature'],
                            max_tokens=DEFAULT_CONFIG['max_tokens'],
                            qdrant_url=os.getenv('QDRANT_URL'),
                            collection_name=PERSISTENT_COLLECTION_NAME,
                            retrieval_k=DEFAULT_CONFIG['retrieval_k'],
                            score_threshold=DEFAULT_CONFIG['score_threshold'],
                            use_custom_llm=False,
                            custom_llm_url=None,
                            custom_llm_api_key=None,
                            custom_llm_model_name=None
                        )
                        st.session_state['chatbot_manager'] = chatbot
                    
                    # Determine RAG usage and layman mode
                    use_rag = True
                    layman_mode = False
                    
                    if chatbot_mode == "General Chat":
                        use_rag = False
                    elif chatbot_mode == "Hybrid (Smart)":
                        use_rag = st.session_state.get('collection_created', False)
                    elif chatbot_mode == "Layman Explanation":
                        use_rag = False
                        layman_mode = True
                    
                    # Get response
                    response = st.session_state['chatbot_manager'].get_response(
                        user_input,
                        enable_content_filter=DEFAULT_CONFIG['enable_content_filter'],
                        enable_pii_detection=DEFAULT_CONFIG['enable_pii_detection'],
                        use_rag=use_rag,
                        layman_mode=layman_mode
                    )
                    
                    tokens_used = response.get('tokens_used', 0)
                    st.session_state.total_tokens_used += tokens_used
                    
                    # Add to history
                    st.session_state['messages'].append({
                        "role": "user", 
                        "content": user_input
                    })
                    st.session_state['messages'].append({
                        "role": "assistant",
                        "content": response.get('answer', 'No response'),
                        "sources": response.get('sources', []),
                        "is_flagged": response.get('is_flagged', False),
                        "flag_reason": response.get('flag_reason'),
                        "tokens_used": tokens_used,
                        "response_type": response.get('response_type', 'rag')
                    })
                    
                except Exception as e:
                    logger.error(f"Chat error: {e}")
                    st.session_state['messages'].append({
                        "role": "assistant",
                        "content": f"⚠️ Error: {str(e)}",
                        "sources": [],
                        "is_flagged": True,
                        "flag_reason": "System error",
                        "tokens_used": 0,
                        "response_type": "error"
                    })
            
            st.rerun()
    else:
        if not st.session_state.get('collection_created') and chatbot_mode == "Document Only":
            st.info("💡 Waiting for embeddings to be created. Please check sidebar for status.")

# Sticky Footer Disclaimer
st.markdown("""
<div class="footer-disclaimer">
    <p>⚠️ <strong>Disclaimer:</strong> This chatbot can make mistakes. Please verify important information from original sources.</p>
</div>
""", unsafe_allow_html=True)

# Cleanup
def cleanup_temp_files():
    for fp in st.session_state.get('temp_file_paths', []):
        if os.path.exists(fp):
            try:
                os.remove(fp)
            except:
                pass

atexit.register(cleanup_temp_files)