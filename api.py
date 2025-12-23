# fastapi.py - Cognitbotz AI Legal Platform - FastAPI Backend
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Depends, Header
import jwt
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
from pathlib import Path
import logging
import os
import tempfile
import uuid
import sqlite3
import glob
from dotenv import load_dotenv
from io import BytesIO

# PDF and DOCX generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load environment
load_dotenv()

# Import our modules
from vectors import EmbeddingsManager
from chatbot import ChatbotManager
from legal_drafting import LegalDraftingManager

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = ['pdf', 'txt', 'docx', 'pptx']
MAX_INTERACTIONS_PER_SESSION = 25
MAX_TOKENS_PER_SESSION = 75000
PERSISTENT_COLLECTION_NAME = "Legal_documents"
DOCUMENTS_FOLDER = Path(__file__).parent / "documents"

# JWT Configuration
SECRET_KEY = os.getenv("JWT_SECRET_KEY", "your-secret-key-change-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 480  # 8 hours

# Temporary storage for anonymous sessions (in-memory)
TEMPORARY_CHAT_STORAGE = {}

# Initialize FastAPI
app = FastAPI(
    title="Cognitbotz AI Legal Platform API",
    description="Backend API for legal document analysis and drafting",
    version="1.0.0"
)

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this properly in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security

# Hardcoded users (same as rag.py)
USERS_DB = {
    "admin": {
        "password": "Admin@123",
        "role": "Administrator",
        "full_name": "System Administrator",
        "user_id": "user_admin_123"
    },
    "finance": {
        "password": "Finance@123",
        "role": "Finance User",
        "full_name": "Finance Department",
        "user_id": "user_finance_456"
    },
    "requester": {
        "password": "Request@123",
        "role": "Purchase Requester",
        "full_name": "Purchase Requester",
        "user_id": "user_requester_789"
    }
}

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    """Create JWT token"""
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


async def verify_token(authorization: Optional[str] = Header(None)) -> Optional[str]:
    """Verify JWT token from auth service and return user email as user_id"""
    if not authorization:
        return None
    
    try:
        scheme, token = authorization.split()
        if scheme.lower() != "bearer":
            return None
        
        # Decode token from auth service (uses 'sub' field for email)
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        
        # Auth service uses 'sub' for email, we use email as user_id
        user_email = payload.get("sub")
        if user_email:
            logger.info(f"[AUTH] Token verified for user: {user_email}")
            return user_email
        
        # Fallback: check for old format with 'user_id' field
        user_id = payload.get("user_id")
        if user_id:
            logger.info(f"[AUTH] Token verified (legacy format) for user_id: {user_id}")
            return user_id
            
        return None
    except Exception as e:
        logger.warning(f"Token validation failed: {e}")
        return None

# In-memory session storage (use Redis in production)
active_sessions = {}

# ==================== PYDANTIC MODELS ====================

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    access_token: str
    token_type: str
    user: Dict[str, str]

class ChatRequest(BaseModel):
    message: str
    chatbot_mode: str = "Document Only"
    chat_session_id: Optional[str] = None
    layman_mode: Optional[bool] = False

class ChatResponse(BaseModel):
    answer: str
    sources: List[Dict[str, Any]]
    tokens_used: int
    is_flagged: bool
    flag_reason: Optional[str]
    response_type: str
    chat_session_id: str

class DraftingRequest(BaseModel):
    doc_type: str
    requirements: str
    style: str = "Formal Legal"
    length: str = "Standard"
    clauses: Optional[List[str]] = None
    special_provisions: Optional[str] = None

class DraftingResponse(BaseModel):
    document: str
    doc_type: str
    style: str
    word_count: int
    tokens_used: int
    metadata: Dict[str, Any]

class StatusResponse(BaseModel):
    embeddings_ready: bool
    collection_name: str
    message: str

class SessionStats(BaseModel):
    interaction_count: int
    total_tokens_used: int
    max_interactions: int
    max_tokens: int
    embeddings_ready: bool

# ==================== DATABASE FUNCTIONS ====================

def init_database():
    """Initialize SQLite database with schema migration support"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        # Create table with current schema
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chat_logs (
                id TEXT PRIMARY KEY,
                username TEXT,
                user_id TEXT,
                chat_session_id TEXT,
                session_id TEXT,
                timestamp TEXT,
                user_input TEXT,
                bot_response TEXT,
                tokens_used INTEGER,
                is_flagged BOOLEAN,
                flag_reason TEXT
            )
        """)
        
        # Migration: Check if user_id column exists
        cursor.execute("PRAGMA table_info(chat_logs)")
        columns = [info[1] for info in cursor.fetchall()]
        
        if 'user_id' not in columns:
            logger.info("Migrating database: Adding user_id column")
            cursor.execute("ALTER TABLE chat_logs ADD COLUMN user_id TEXT")
        
        # Create index for efficient user history retrieval
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_user_id ON chat_logs(user_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_chat_session_id ON chat_logs(chat_session_id)")
        
        conn.commit()
        conn.close()
        logger.info("Database initialized and verified")
    except Exception as e:
        logger.error(f"Database init error: {e}")

def migrate_old_conversations():
    """Migrate old conversations to associate them with users by username"""
    try:
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        # Check total conversations in database
        cursor.execute("SELECT COUNT(*) FROM chat_logs")
        total_count = cursor.fetchone()[0]
        logger.info(f"[MIGRATION] Total conversations in database: {total_count}")
        
        # Find all conversations with NULL user_id
        cursor.execute("""
            SELECT DISTINCT username, COUNT(*) as count FROM chat_logs 
            WHERE user_id IS NULL AND username != 'anonymous'
            GROUP BY username
        """)
        
        old_conversations = cursor.fetchall()
        logger.info(f"[MIGRATION] Conversations with NULL user_id: {old_conversations}")
        
        if not old_conversations:
            logger.info("[MIGRATION] No old conversations to migrate")
            # Check how many have user_id already set
            cursor.execute("SELECT COUNT(*) FROM chat_logs WHERE user_id IS NOT NULL")
            migrated_count = cursor.fetchone()[0]
            logger.info(f"[MIGRATION] Conversations with user_id already set: {migrated_count}")
            conn.close()
            return
        
        logger.info(f"[MIGRATION] Found {len(old_conversations)} users with old conversations to migrate")
        
        # Map usernames to user_ids from USERS_DB
        total_migrated = 0
        for username, count in old_conversations:
            logger.info(f"[MIGRATION] Processing user: {username} with {count} conversations")
            if username in USERS_DB:
                user_id = USERS_DB[username].get('user_id')
                logger.info(f"[MIGRATION] Mapped {username} to user_id: {user_id}")
                if user_id:
                    # Update all conversations for this user
                    cursor.execute("""
                        UPDATE chat_logs 
                        SET user_id = ? 
                        WHERE username = ? AND user_id IS NULL
                    """, (user_id, username))
                    
                    migrated_count = cursor.rowcount
                    total_migrated += migrated_count
                    logger.info(f"[MIGRATION] Migrated {migrated_count} conversations for user {username}")
            else:
                logger.warning(f"[MIGRATION] User {username} not found in USERS_DB")
        
        conn.commit()
        
        # Verify migration
        cursor.execute("SELECT COUNT(*) FROM chat_logs WHERE user_id IS NULL AND username != 'anonymous'")
        remaining = cursor.fetchone()[0]
        cursor.execute("SELECT COUNT(*) FROM chat_logs WHERE user_id IS NOT NULL")
        with_user_id = cursor.fetchone()[0]
        
        logger.info(f"[MIGRATION] Completed: {total_migrated} total migrated")
        logger.info(f"[MIGRATION] Remaining NULL user_id conversations (non-anonymous): {remaining}")
        logger.info(f"[MIGRATION] Total conversations with user_id: {with_user_id}")
        
        conn.close()
    except Exception as e:
        logger.error(f"[MIGRATION] Error: {e}", exc_info=True)

def log_interaction(username: str, chat_session_id: str, user_input: str, 
                   bot_response: str, tokens_used: int, is_flagged: bool = False, 
                   flag_reason: str = None, user_id: str = None):
    """Log chat interaction - only for authenticated users"""
    try:
        # Only persist chat for authenticated users, not anonymous users
        if user_id is not None:
            logger.info(f"[LOG_INTERACTION] SAVING: user_id={user_id}, chat_session_id={chat_session_id}, username={username}")
            logger.info(f"[LOG_INTERACTION] Message preview: {user_input[:60]}...")
            
            conn = sqlite3.connect("chat_logs.db", timeout=10)
            cursor = conn.cursor()
            
            log_id = str(uuid.uuid4())
            timestamp = datetime.now().isoformat()
            
            cursor.execute("""
                INSERT INTO chat_logs 
                (id, username, user_id, chat_session_id, session_id, timestamp, user_input, 
                 bot_response, tokens_used, is_flagged, flag_reason)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (log_id, username or "anonymous", user_id, chat_session_id, str(uuid.uuid4()), 
                  timestamp, user_input, bot_response, tokens_used, is_flagged, flag_reason))
            
            conn.commit()
            
            # Verify it was saved
            cursor.execute("SELECT COUNT(*) FROM chat_logs WHERE chat_session_id = ? AND user_id = ?", 
                          (chat_session_id, user_id))
            count = cursor.fetchone()[0]
            logger.info(f"[LOG_INTERACTION] VERIFIED: Found {count} messages for this session with this user_id")
            
            conn.close()
            logger.info(f"[LOG_INTERACTION] SUCCESS: Message saved to database")
        else:
            logger.info(f"[LOG_INTERACTION] ANONYMOUS: Storing in memory (user_id is None)")
            # For anonymous users, store in temporary memory storage
            # Use the session_key to group temporary chats
            session_key = f"anonymous_{user_id or 'temp'}" if user_id else f"anonymous_{str(uuid.uuid4())}"
            if session_key not in TEMPORARY_CHAT_STORAGE:
                TEMPORARY_CHAT_STORAGE[session_key] = {}
                    
            if chat_session_id not in TEMPORARY_CHAT_STORAGE[session_key]:
                TEMPORARY_CHAT_STORAGE[session_key][chat_session_id] = []
                    
            timestamp = datetime.now().isoformat()
            TEMPORARY_CHAT_STORAGE[session_key][chat_session_id].append({
                "timestamp": timestamp,
                "user_input": user_input,
                "bot_response": bot_response,
                "tokens_used": tokens_used,
                "is_flagged": is_flagged,
                "flag_reason": flag_reason
            })
    except Exception as e:
        logger.error(f"[LOG_INTERACTION] FAILED: {e}", exc_info=True)

def get_user_chat_sessions(user_identifier: str):
    """Get all chat sessions for a user (by user_id which is now email)"""
    try:
        sessions = []
        
        # If user_identifier is present and not empty, treat as authenticated user (email)
        # Otherwise, treat as anonymous
        if user_identifier:
            # Authenticated user: Query by user_id (email) from persistent storage
            conn = sqlite3.connect("chat_logs.db", timeout=10)
            cursor = conn.cursor()
            
            logger.info(f"[QUERY] Getting chat sessions for user_id: {user_identifier}")
            
            # First check what's in the database
            cursor.execute("SELECT COUNT(*) FROM chat_logs WHERE user_id = ?", (user_identifier,))
            total_for_user = cursor.fetchone()[0]
            logger.info(f"[QUERY] Total messages in database for this user_id: {total_for_user}")
            
            cursor.execute("SELECT COUNT(*) FROM chat_logs")
            total_in_db = cursor.fetchone()[0]
            logger.info(f"[QUERY] Total messages in entire database: {total_in_db}")
            
            query = """
                SELECT 
                    chat_session_id,
                    MIN(timestamp) as first_msg_time,
                    COUNT(*) as message_count,
                    MAX(user_input) as first_input
                FROM chat_logs
                WHERE user_id = ? AND chat_session_id IS NOT NULL
                GROUP BY chat_session_id
                ORDER BY first_msg_time DESC
            """
            
            cursor.execute(query, (user_identifier,))
            rows = cursor.fetchall()
            logger.info(f"[QUERY] Found {len(rows)} chat sessions for user_id {user_identifier}")
            conn.close()
            
            for row in rows:
                chat_id, timestamp, count, first_input = row
                try:
                    dt = datetime.fromisoformat(timestamp)
                    date_str = dt.strftime("%b %d, %Y %I:%M %p")
                except:
                    date_str = timestamp
                
                preview = (first_input[:40] + "...") if first_input and len(first_input) > 40 else (first_input or "Empty chat")
                logger.info(f"[QUERY] Chat session {chat_id}: {count} messages")
                
                sessions.append({
                    "id": chat_id,
                    "date": date_str,
                    "count": count,
                    "preview": preview
                })
        else:
            # Anonymous user: Return temporary session IDs from memory
            session_key = f"anonymous_{user_identifier or 'temp'}" if user_identifier else f"anonymous_{str(uuid.uuid4())}"
            if session_key in TEMPORARY_CHAT_STORAGE:
                for chat_session_id, messages in TEMPORARY_CHAT_STORAGE[session_key].items():
                    if messages:  # If session has messages
                        first_message = messages[0]
                        timestamp = first_message["timestamp"]
                        first_input = first_message["user_input"]
                        
                        try:
                            dt = datetime.fromisoformat(timestamp)
                            date_str = dt.strftime("%b %d, %Y %I:%M %p")
                        except:
                            date_str = timestamp
                        
                        preview = (first_input[:40] + "...") if first_input and len(first_input) > 40 else (first_input or "Empty chat")
                        
                        sessions.append({
                            "id": chat_session_id,
                            "date": date_str,
                            "count": len(messages),
                            "preview": preview
                        })
        
        return sessions
    except Exception as e:
        logger.error(f"Error loading chat sessions: {e}")
        return []

def get_chat_messages(chat_session_id: str):
    """Get all messages from a specific chat session"""
    try:
        messages = []
        
        # First, check if this session exists in any temporary storage
        found_in_temp = False
        for session_key, session_data in TEMPORARY_CHAT_STORAGE.items():
            if chat_session_id in session_data:
                # Load from temporary memory storage
                session_messages = session_data[chat_session_id]
                for msg in session_messages:
                    messages.append({
                        "role": "user",
                        "content": msg["user_input"]
                    })
                    messages.append({
                        "role": "assistant",
                        "content": msg["bot_response"],
                        "tokens_used": msg["tokens_used"],
                        "is_flagged": msg["is_flagged"],
                        "flag_reason": msg["flag_reason"]
                    })
                found_in_temp = True
                break
        
        if not found_in_temp:
            # Load from persistent database (for authenticated users)
            conn = sqlite3.connect("chat_logs.db", timeout=10)
            cursor = conn.cursor()
            
            logger.info(f"[MESSAGES] Loading messages for chat_session_id: {chat_session_id}")
            cursor.execute("""
                SELECT user_input, bot_response, tokens_used, is_flagged, flag_reason
                FROM chat_logs
                WHERE chat_session_id = ?
                ORDER BY timestamp ASC
            """, (chat_session_id,))
            
            rows = cursor.fetchall()
            logger.info(f"[MESSAGES] Found {len(rows)} messages for chat_session_id {chat_session_id}")
            conn.close()
            
            for row in rows:
                user_input, bot_response, tokens, flagged, flag_reason = row
                
                messages.append({
                    "role": "user",
                    "content": user_input
                })
                messages.append({
                    "role": "assistant",
                    "content": bot_response,
                    "tokens_used": tokens,
                    "is_flagged": bool(flagged),
                    "flag_reason": flag_reason
                })
        
        return messages
    except Exception as e:
        logger.error(f"Error loading chat messages: {e}")
        return []

def delete_chat_session_db(chat_session_id: str, user_identifier: str):
    """Delete a chat session"""
    try:
        # Check if this is a temporary session (anonymous user)
        session_key = f"anonymous_{user_identifier or 'temp'}" if user_identifier else f"anonymous_{str(uuid.uuid4())}"
        if session_key in TEMPORARY_CHAT_STORAGE and chat_session_id in TEMPORARY_CHAT_STORAGE[session_key]:
            # Delete from temporary memory storage
            del TEMPORARY_CHAT_STORAGE[session_key][chat_session_id]
            # If the session_key has no more chat sessions, remove it completely
            if not TEMPORARY_CHAT_STORAGE[session_key]:
                del TEMPORARY_CHAT_STORAGE[session_key]
            return True
        else:
            # Delete from persistent database (authenticated user)
            conn = sqlite3.connect("chat_logs.db", timeout=10)
            cursor = conn.cursor()
            
            # Determine if we are deleting for an authenticated user (email) or anonymous
            logger.info(f"[DELETE] Attempting to delete chat_session_id: {chat_session_id}, user_identifier: {user_identifier}")
                
            if user_identifier:
                # Authenticated user - delete by user_id (email)
                cursor.execute("""
                    SELECT COUNT(*) FROM chat_logs 
                    WHERE chat_session_id = ? AND user_id = ?
                """, (chat_session_id, user_identifier))
                count = cursor.fetchone()[0]
                logger.info(f"[DELETE] Found {count} records matching user_id")
                
                # If not found by user_id, try without user_id filter (for migration purposes)
                if count == 0:
                    cursor.execute("""
                        SELECT COUNT(*) FROM chat_logs 
                        WHERE chat_session_id = ?
                    """, (chat_session_id,))
                    count = cursor.fetchone()[0]
                    logger.info(f"[DELETE] Found {count} records without user_id filter")
                    
                    # Delete without user_id filter if found
                    if count > 0:
                        cursor.execute("""
                            DELETE FROM chat_logs 
                            WHERE chat_session_id = ?
                        """, (chat_session_id,))
                else:
                    # Delete with user_id filter
                    cursor.execute("""
                        DELETE FROM chat_logs 
                        WHERE chat_session_id = ? AND user_id = ?
                    """, (chat_session_id, user_identifier))
            else:
                # Anonymous - delete by chat_session_id only
                cursor.execute("""
                    DELETE FROM chat_logs 
                    WHERE chat_session_id = ? AND user_id IS NULL
                """, (chat_session_id,))
            
            conn.commit()
            deleted = cursor.rowcount
            logger.info(f"[DELETE] Deleted {deleted} records")
            conn.close()
            
            return deleted > 0
        
        return True  # For temporary sessions, we always return True
    except Exception as e:
        logger.error(f"Error deleting chat session: {e}")
        return False

# ==================== AUTO-EMBEDDING FUNCTIONS ====================

def check_collection_exists():
    """Check if embeddings collection exists"""
    try:
        from qdrant_client import QdrantClient
        
        qdrant_url = os.getenv('QDRANT_URL')
        qdrant_api_key = os.getenv('QDRANT_API_KEY')
        
        if not qdrant_url or not qdrant_api_key:
            return False
        
        client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key, prefer_grpc=False)
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]
        
        if PERSISTENT_COLLECTION_NAME in collection_names:
            collection_info = client.get_collection(PERSISTENT_COLLECTION_NAME)
            return collection_info.points_count > 0
        
        return False
    except Exception as e:
        logger.error(f"Error checking collection: {e}")
        return False

def get_documents_from_folder():
    """Get all documents from static folder"""
    if not DOCUMENTS_FOLDER.exists():
        return []
    
    document_files = []
    for ext in ALLOWED_EXTENSIONS:
        pattern = str(DOCUMENTS_FOLDER / f"*.{ext}")
        files = glob.glob(pattern)
        document_files.extend(files)
    
    return document_files

def auto_embed_documents():
    """Auto-embed documents on startup"""
    try:
        if check_collection_exists():
            return True, "Embeddings already exist"
        
        document_files = get_documents_from_folder()
        if not document_files:
            return False, "No documents found"
        
        embeddings_mgr = EmbeddingsManager(
            model_name="BAAI/bge-small-en",
            device="cpu",
            encode_kwargs={"normalize_embeddings": True},
            qdrant_url=os.getenv('QDRANT_URL'),
            collection_name=PERSISTENT_COLLECTION_NAME,
            chunk_size=1200,
            chunk_overlap=300
        )
        
        processed = 0
        for doc_path in document_files:
            try:
                embeddings_mgr.create_embeddings(doc_path)
                processed += 1
            except Exception as e:
                logger.error(f"Failed to process {doc_path}: {e}")
        
        if processed > 0:
            return True, f"Auto-embedded {processed} documents"
        return False, "Failed to embed documents"
    except Exception as e:
        logger.error(f"Auto-embedding error: {e}")
        return False, str(e)

# ==================== DOCUMENT GENERATION ====================

def generate_pdf_document(content: str, doc_type: str, metadata: dict) -> BytesIO:
    """Generate PDF document"""
    if not REPORTLAB_AVAILABLE:
        return None
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='#1e3a8a',
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=12,
        fontName='Times-Roman',
        leading=14
    )
    
    story = []
    story.append(Paragraph(doc_type.upper(), title_style))
    story.append(Spacer(1, 0.2*inch))
    
    metadata_text = f"<b>Generated:</b> {metadata.get('generated_at', 'N/A')}<br/>"
    metadata_text += f"<b>Jurisdiction:</b> {metadata.get('jurisdiction', 'General')}"
    story.append(Paragraph(metadata_text, body_style))
    story.append(Spacer(1, 0.3*inch))
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            story.append(Paragraph(line, body_style))
    
    footer_text = "<i>Generated by Cognitbotz AI Legal Platform</i>"
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(footer_text, ParagraphStyle(
        'Footer', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER
    )))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_document(content: str, doc_type: str, metadata: dict) -> BytesIO:
    """Generate DOCX document"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = DocxDocument()
    
    doc.core_properties.title = doc_type
    doc.core_properties.author = "Cognitbotz AI Legal Platform"
    
    title = doc.add_heading(doc_type.upper(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.add_run('Generated: ').bold = True
    meta_para.add_run(metadata.get('generated_at', 'N/A'))
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run('Generated by Cognitbotz AI Legal Platform')
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== JWT & AUTH ====================

def get_session_manager(user_id: str = None):
    """Get or create session manager for user or anonymous"""
    # For authenticated users, use user_id as session key
    # For anonymous users, generate a temporary session key
    if user_id:
        session_key = user_id
    else:
        # Create a temporary session key for anonymous users
        # Use a consistent key for the same anonymous session
        session_key = f"anonymous_{str(uuid.uuid4())}"
    
    if session_key not in active_sessions:
        active_sessions[session_key] = {
            "session_id": str(uuid.uuid4()),
            "chat_session_id": str(uuid.uuid4()),
            "chatbot_manager": None,
            "drafting_manager": None,
            "interaction_count": 0,
            "total_tokens_used": 0,
            "created_at": datetime.now(),
            "user_id": user_id
        }
    return active_sessions[session_key]

def get_or_create_anonymous_session():
    """Legacy wrapper for anonymous session"""
    return get_session_manager(None)

# ==================== STARTUP ====================

@app.on_event("startup")
async def startup_event():
    """Initialize on startup"""
    logger.info("Starting Cognitbotz AI Legal Platform API...")
    
    # Initialize database
    init_database()
    
    # Migrate old conversations to add user_id
    migrate_old_conversations()
    
    # Auto-embed documents
    logger.info("Checking for embeddings...")
    success, message = auto_embed_documents()
    if success:
        logger.info(f"✅ {message}")
    else:
        logger.warning(f"⚠️ {message}")
    
    # Initialize anonymous session
    get_or_create_anonymous_session()

# ==================== API ENDPOINTS ====================

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "service": "Cognitbotz AI Legal Platform API",
        "version": "1.0.0",
        "status": "running"
    }

# Authentication is handled by authentication/auth_service.py
# This API only consumes JWT tokens from that service

@app.get("/api/status", response_model=StatusResponse)
async def get_status():
    """Get embeddings status"""
    embeddings_ready = check_collection_exists()
    
    return StatusResponse(
        embeddings_ready=embeddings_ready,
        collection_name=PERSISTENT_COLLECTION_NAME,
        message="Embeddings ready" if embeddings_ready else "Embeddings not ready"
    )

@app.get("/api/session/stats", response_model=SessionStats)
async def get_session_stats():
    """Get session statistics"""
    session = get_or_create_anonymous_session()
    embeddings_ready = check_collection_exists()
    
    return SessionStats(
        interaction_count=session["interaction_count"],
        total_tokens_used=session["total_tokens_used"],
        max_interactions=MAX_INTERACTIONS_PER_SESSION,
        max_tokens=MAX_TOKENS_PER_SESSION,
        embeddings_ready=embeddings_ready
    )

@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, user_id: Optional[str] = Depends(verify_token)):
    """Send chat message and get response"""
    session = get_session_manager(user_id)
    
    # Check limits
    if session["interaction_count"] >= MAX_INTERACTIONS_PER_SESSION:
        raise HTTPException(status_code=429, detail="Interaction limit reached")
    
    if session["total_tokens_used"] >= MAX_TOKENS_PER_SESSION:
        raise HTTPException(status_code=429, detail="Token limit reached")
    
    # Use provided chat_session_id or current one
    chat_session_id = request.chat_session_id or session["chat_session_id"]
    
    try:
        # Initialize chatbot if needed
        if not session["chatbot_manager"]:
            session["chatbot_manager"] = ChatbotManager(
                model_name="BAAI/bge-small-en",
                device="cpu",
                encode_kwargs={"normalize_embeddings": True},
                llm_model=os.getenv("LLM_MODEL", "llama-3.3-70b-versatile"),
                llm_temperature=0.3,
                max_tokens=3000,
                qdrant_url=os.getenv('QDRANT_URL'),
                collection_name=PERSISTENT_COLLECTION_NAME,
                retrieval_k=5,
                score_threshold=0.5,
                use_custom_llm=False,
                custom_llm_url=None,
                custom_llm_api_key=None,
                custom_llm_model_name=None
            )
        
        # Determine RAG usage and layman mode
        use_rag = True
        layman_mode = request.layman_mode or False
        
        if request.chatbot_mode == "General Chat":
            use_rag = False
        elif request.chatbot_mode == "Hybrid (Smart)":
            use_rag = check_collection_exists()
        elif request.chatbot_mode == "Layman Explanation":
            use_rag = False
            layman_mode = True
        
        # Get response
        response = session["chatbot_manager"].get_response(
            request.message,
            enable_content_filter=True,
            enable_pii_detection=True,
            use_rag=use_rag,
            layman_mode=layman_mode
        )
        
        tokens_used = response.get('tokens_used', 0)
        session["interaction_count"] += 1
        session["total_tokens_used"] += tokens_used
        
        # Log interaction
        # user_id from verify_token is now the email from auth service
        username_to_log = user_id if user_id else 'anonymous'
        
        log_interaction(
            username=username_to_log,
            user_id=user_id,
            chat_session_id=chat_session_id,
            user_input=request.message,
            bot_response=response.get('answer', ''),
            tokens_used=tokens_used,
            is_flagged=response.get('is_flagged', False),
            flag_reason=response.get('flag_reason')
        )
        
        return ChatResponse(
            answer=response.get('answer', 'No response'),
            sources=response.get('sources', []),
            tokens_used=tokens_used,
            is_flagged=response.get('is_flagged', False),
            flag_reason=response.get('flag_reason'),
            response_type=response.get('response_type', 'rag'),
            chat_session_id=chat_session_id
        )
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/chat/history")
async def get_chat_history(user_id: Optional[str] = Depends(verify_token)):
    """Get user's chat history"""
    # If authenticated, get user specific sessions from persistent storage
    # If anonymous (user_id is None), get temporary sessions from memory
    if user_id:
        # Authenticated user - query persistent database
        sessions = get_user_chat_sessions(user_id)
    else:
        # Anonymous user - return empty list (temporary chats are in memory)
        sessions = []
    
    return {"sessions": sessions}

@app.get("/api/chat/{chat_session_id}")
async def get_chat(chat_session_id: str, user_id: Optional[str] = Depends(verify_token)):
    """Get specific chat session messages - allows access if conversation exists"""
    try:
        logger.info(f"[GET_CHAT] Request received for chat_session_id: {chat_session_id}, user_id: {user_id}")
        
        conn = sqlite3.connect("chat_logs.db", timeout=10)
        cursor = conn.cursor()
        
        # Check if conversation exists (either with user_id or without)
        if user_id:
            logger.info(f"[GET_CHAT] Authenticated user - attempting to get chat_session_id: {chat_session_id}")
            
            # Try with user_id filter first
            cursor.execute("""
                SELECT COUNT(*) FROM chat_logs 
                WHERE chat_session_id = ? AND user_id = ?
            """, (chat_session_id, user_id))
            count_with_user = cursor.fetchone()[0]
            logger.info(f"[GET_CHAT] Found {count_with_user} records with user_id filter")
            
            # Fallback: try without user_id filter
            if count_with_user == 0:
                cursor.execute("""
                    SELECT COUNT(*) FROM chat_logs 
                    WHERE chat_session_id = ?
                """, (chat_session_id,))
                count_without_user = cursor.fetchone()[0]
                logger.info(f"[GET_CHAT] Found {count_without_user} records without user_id filter (FALLBACK)")
                
                if count_without_user > 0:
                    logger.info(f"[GET_CHAT] Using fallback - session exists but may not have user_id")
                else:
                    logger.error(f"[GET_CHAT] Session not found in database at all")
                    conn.close()
                    raise HTTPException(status_code=404, detail="Chat session not found")
            else:
                logger.info(f"[GET_CHAT] Session found with user_id match")
        else:
            logger.info(f"[GET_CHAT] Anonymous user - checking temporary storage")
            found = False
            for session_key, session_data in TEMPORARY_CHAT_STORAGE.items():
                if chat_session_id in session_data:
                    found = True
                    logger.info(f"[GET_CHAT] Found session in temporary storage: {session_key}")
                    break
            
            if not found:
                logger.error(f"[GET_CHAT] Session not found in temporary storage")
                conn.close()
                raise HTTPException(status_code=404, detail="Chat session not found")
        
        conn.close()
        logger.info(f"[GET_CHAT] Loading messages for session")
        messages = get_chat_messages(chat_session_id)
        logger.info(f"[GET_CHAT] Successfully loaded {len(messages)} messages")
        return {"messages": messages, "chat_session_id": chat_session_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[GET_CHAT] Error retrieving chat: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/chat/{chat_session_id}")
async def delete_chat(chat_session_id: str, user_id: Optional[str] = Depends(verify_token)):
    """Delete a chat session"""
    logger.info(f"[DELETE_ENDPOINT] Request to delete chat_session_id: {chat_session_id}, user_id: {user_id}")
    target_user = user_id if user_id else "anonymous"
    logger.info(f"[DELETE_ENDPOINT] Calling delete_chat_session_db with target_user: {target_user}")
    success = delete_chat_session_db(chat_session_id, target_user)
    logger.info(f"[DELETE_ENDPOINT] Delete result: {success}")
    if success:
        return {"message": "Chat deleted successfully"}
    logger.error(f"[DELETE_ENDPOINT] Failed to delete chat_session_id: {chat_session_id}")
    raise HTTPException(status_code=404, detail="Chat not found")

@app.post("/api/chat/new")
async def new_chat(user_id: Optional[str] = Depends(verify_token)):
    """Start a new chat session"""
    session = get_session_manager(user_id)
    new_chat_id = str(uuid.uuid4())
    session["chat_session_id"] = new_chat_id
    session["interaction_count"] = 0
    
    return {"chat_session_id": new_chat_id, "message": "New chat started"}

@app.post("/api/documents/upload")
async def upload_documents(
    files: List[UploadFile] = File(...)
):
    """Upload and process documents
    
    Note: Documents in the /documents folder are auto-embedded on startup.
    This endpoint is for adding additional documents dynamically.
    """
    processed = []
    failed = []
    
    try:
        embeddings_mgr = EmbeddingsManager(
            model_name="BAAI/bge-small-en",
            device="cpu",
            encode_kwargs={"normalize_embeddings": True},
            qdrant_url=os.getenv('QDRANT_URL'),
            collection_name=PERSISTENT_COLLECTION_NAME,
            chunk_size=1200,
            chunk_overlap=300
        )
        
        for file in files:
            # Validate file
            if file.size > MAX_FILE_SIZE:
                failed.append({"file": file.filename, "reason": "File too large"})
                continue
            
            ext = file.filename.split('.')[-1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                failed.append({"file": file.filename, "reason": "Invalid file type"})
                continue
            
            # Save temporarily and process
            try:
                temp_dir = tempfile.mkdtemp()
                temp_path = os.path.join(temp_dir, file.filename)
                
                with open(temp_path, 'wb') as f:
                    content = await file.read()
                    f.write(content)
                
                embeddings_mgr.create_embeddings(temp_path)
                processed.append(file.filename)
                
                # Cleanup
                os.remove(temp_path)
                os.rmdir(temp_dir)
                
            except Exception as e:
                logger.error(f"Failed to process {file.filename}: {e}")
                failed.append({"file": file.filename, "reason": str(e)})
        
        return {
            "processed": processed,
            "failed": failed,
            "message": f"Processed {len(processed)} documents"
        }
        
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/documents/process")
async def documents_process_double_slash(files: List[UploadFile] = File(...)):
    """Handle POST request for double slash documents process endpoint"""
    return await upload_documents(files)

# Add the missing /api/documents/process endpoint
@app.post("/api/documents/process")
async def documents_process(files: List[UploadFile] = File(...)):
    """Process uploaded documents"""
    return await upload_documents(files)

@app.post("/api/drafting/generate", response_model=DraftingResponse)
async def generate_document(
    request: DraftingRequest,
    user_id: Optional[str] = Depends(verify_token)
):
    """Generate legal document"""
    session = get_session_manager(user_id)
    
    try:
        # Initialize drafting manager if needed
        if not session["drafting_manager"]:
            session["drafting_manager"] = LegalDraftingManager(
                llm_model=os.getenv("LLM_MODEL", "llama-3.3-70b-versatile"),
                temperature=0.3,
                max_tokens=12000,
                user_preferences={}
            )
        
        # Generate document
        result = session["drafting_manager"].generate_document(
            doc_type=request.doc_type,
            prompt=request.requirements,
            style=request.style,
            length=request.length,
            clauses=request.clauses,
            special_provisions=request.special_provisions or ""
        )
        
        if 'error' in result:
            raise HTTPException(status_code=500, detail=result['error'])
        
        tokens_used = result.get('tokens_used', 0)
        session["total_tokens_used"] += tokens_used
        
        # Log interaction
        username_to_log = None
        if user_id:
            # Find the username by matching user_id in USERS_DB
            for user_name, user_data in USERS_DB.items():
                if user_data.get('user_id') == user_id:
                    username_to_log = user_name
                    break
        else:
            username_to_log = 'anonymous'
        
        log_interaction(
            username=username_to_log,
            user_id=user_id,
            chat_session_id=session["chat_session_id"],
            user_input=request.requirements,
            bot_response=result.get('document', '')[:1000],
            tokens_used=tokens_used
        )
        return DraftingResponse(
            document=result.get('document', ''),
            doc_type=result.get('doc_type', ''),
            style=result.get('style', ''),
            word_count=result.get('word_count', 0),
            tokens_used=tokens_used,
            metadata=result.get('metadata', {})
        )
        
    except Exception as e:
        logger.error(f"Drafting error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/drafting/download/{format}")
async def download_document(
    format: str,
    content: str = Form(...),
    doc_type: str = Form(...)
):
    """Download generated document in specified format"""
    
    metadata = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "jurisdiction": "General"
    }
    
    if format == "txt":
        return StreamingResponse(
            iter([content]),
            media_type="text/plain",
            headers={
                "Content-Disposition": f"attachment; filename={doc_type.replace(' ', '_')}.txt"
            }
        )
    
    elif format == "pdf":
        if not REPORTLAB_AVAILABLE:
            raise HTTPException(status_code=400, detail="PDF generation not available")
        
        pdf_buffer = generate_pdf_document(content, doc_type, metadata)
        if not pdf_buffer:
            raise HTTPException(status_code=500, detail="PDF generation failed")
        
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={doc_type.replace(' ', '_')}.pdf"
            }
        )
    
    elif format == "docx":
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=400, detail="DOCX generation not available")
        
        docx_buffer = generate_docx_document(content, doc_type, metadata)
        if not docx_buffer:
            raise HTTPException(status_code=500, detail="DOCX generation failed")
        
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={doc_type.replace(' ', '_')}.docx"
            }
        )
    
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

@app.post("/api/session/reset")
async def reset_session(user_id: Optional[str] = Depends(verify_token)):
    """Reset user session"""
    # For authenticated users, clear their session
    if user_id and user_id in active_sessions:
        del active_sessions[user_id]
    # For anonymous users, clear their temporary session
    elif user_id is None:
        # Find and remove anonymous session
        keys_to_remove = []
        for key in active_sessions:
            if key.startswith("anonymous_"):
                keys_to_remove.append(key)
        for key in keys_to_remove:
            if key in active_sessions:
                del active_sessions[key]
        
        # Create a new anonymous session
        get_or_create_anonymous_session()
    
    return {"message": "Session reset successfully"}

@app.post("/api/auth/logout")
async def logout(user_id: Optional[str] = Depends(verify_token)):
    """Logout user"""
    # For authenticated users, clear their session
    if user_id and user_id in active_sessions:
        del active_sessions[user_id]
    # For anonymous users, clear their temporary session
    elif user_id is None:
        # Find and remove anonymous session
        keys_to_remove = []
        for key in active_sessions:
            if key.startswith("anonymous_"):
                keys_to_remove.append(key)
        for key in keys_to_remove:
            if key in active_sessions:
                del active_sessions[key]
    
    return {"message": "Logged out successfully"}

# ==================== HEALTH CHECK ====================

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    embeddings_ready = check_collection_exists()
    
    return {
        "status": "healthy",
        "embeddings_ready": embeddings_ready,
        "collection": PERSISTENT_COLLECTION_NAME,
        "timestamp": datetime.now().isoformat()
    }

# ==================== RUN SERVER ====================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")